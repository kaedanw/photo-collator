import customtkinter as ctk # pip install customtkinter | https://customtkinter.tomschimansky.com/
from customtkinter import filedialog
# from CTkMessagebox import CTkMessagebox # pip install CTkMessagebox | https://github.com/Akascape/CTkMessagebox
import os
from docx import Document # pip install python-docx | https://python-docx.readthedocs.io/en/latest/user/install.html
from docx.shared import Pt, Cm, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import natsort # pip install natsort | https://pypi.org/project/natsort/#installation
import threading
from PIL import Image

# tkinter mainloop is blocking. handle frozen gui with threading
def threaded_func(func):
    def inner(*args, **kwargs):
        thread = threading.Thread(target=func, args=args, kwargs=kwargs)
        thread.start()
    return inner

@threaded_func
def create_word_document(folder_paths, pic_width, pic_height, start_count, default_caption, pic_extensions, save_file):
    # Get picture dimensions
    pic_width = float(pic_width)
    pic_height = float(pic_height)

    # Choose Word document
    mode = app_mode.get()
    if mode == menu_values[0]: # Create New Document
        print("Creating Word Document..")
        doc = Document()
    else: # Load Existing Document
        print("Loading Word Document..")
        doc = Document(save_file)

        if mode == menu_values[2]: # Insert Mode - Between Photos
            # Find start count photo to insert_before_paragraph
            paragraph_before = None
            for i, paragraph in enumerate(doc.paragraphs):
                if paragraph.text:
                    if paragraph.text.split()[-1] == start_count:
                        paragraph_before = doc.paragraphs[i-1]
                        break
            if not paragraph_before:
                print("Existing photo to insert new photos before could not be found..")
                start_count_entry.configure(border_color="pink")
                status_label.configure(text="ERROR: Existing photo could not be found..", fg_color="pink")
                return
    
    # Check file is not opened:
    try:
        doc.save(save_file)
    except PermissionError:
        return status_label.configure(text="ERROR: File is opened. Please close and retry.", fg_color="pink")
    
    # Set landscape orientation and two columns layout
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Mm(17.5)
    section.right_margin = Mm(17.5)
    section.top_margin = Mm(12.5)
    section.bottom_margin = Mm(10.0)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')
    
    # Set document style
    style = doc.styles['Normal']
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.1

    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Initialize photo counter
    if mode == menu_values[1]: # Append Mode - End of document
        count = int(doc.paragraphs[-1].text.split()[-1]) + 1
    else:
        count = int(start_count)

    # Iterate through each folder entered
    for folder_path in folder_paths:
        # Get a list of photos with the specified extension in the folder
        pic_files = [f for f in os.listdir(folder_path) if any(f.upper().endswith(e) for e in pic_extensions)]
        pic_files = natsort.natsorted(pic_files, alg=natsort.PATH)

        # Iterate through the picture files and insert them into the document
        for pic_file in pic_files:
            img_path = os.path.join(folder_path, pic_file)
            rotated = process_image(img_path)
            caption = f"{default_caption} {count}"
            if mode == menu_values[2]:
                paragraph_before.insert_paragraph_before().add_run().add_picture(img_path, width=Cm(pic_width), height=Cm(pic_height))
                paragraph_before.insert_paragraph_before().add_run(caption)
            else:
                doc.add_paragraph()
                doc.paragraphs[-1].add_run().add_picture(img_path, width=Cm(pic_width), height=Cm(pic_height))
                
                # Add a caption below the image
                doc.add_paragraph()
                doc.paragraphs[-1].add_run(caption)
            
            if rotated:
                process_image(img_path, 90)
            status_label.configure(text=f"Inserting Photo {count}..", fg_color="yellow")
            print(f"Inserting Photo {count}.. [{pic_file}]")
            count += 1

    # In insert mode, rename sequential photos
    if mode == menu_values[2]:
        rename = False
        for paragraph in doc.paragraphs:
            if paragraph.text:
                num = paragraph.text.split()[-1]
                if rename:
                    paragraph.text = paragraph.text.replace(num, str(count))
                    count += 1
                else:
                    if num == str(count - 1):
                        rename = True

    # Save the Word document
    if count == int(start_count):
        print("No photos with the specified extension was found..")
        status_label.configure(text="FINISHED: No photos found..", fg_color="lightgreen")
        return
    status_label.configure(text=f"Saving document.. Please Wait.")
    print("Saving document.. Please Wait.")
    doc.save(save_file)
    print("Word document collated successfully.")
    status_label.configure(text="FINISHED SUCCESSFULLY", fg_color="lightgreen")

def process_image(img_path, rotation_angle=-90):
    rotated_flag = False
    image = Image.open(img_path)
    # Ensure picture is landscape in the document, revert after adding
    if (rotation_angle == 90) or (image.width < image.height):
        image = image.rotate(rotation_angle, expand=True)
        image.save(img_path)
        rotated_flag = True
    image.close()
    return rotated_flag

def browse_folder():
    folder_path = filedialog.askdirectory()
    if not folder_path:
        return
    entry_lines = get_folder_entry()
    if entry_lines != ['']:
        folder_path_entry.insert("end", '\n')
    folder_path_entry.insert("end", folder_path)
    folder_check()

def browse_save():
    if app_mode.get() == menu_values[0]: # Create Mode - Create Document
        # (Optional): Use desktop as default file path on file dialog
        # desktop = os.path.normpath(os.path.expanduser("~/Desktop")).replace('\\','/')
        # save_file =  filedialog.asksaveasfilename(initialdir=desktop, initialfile="Photos.docx", defaultextension=".docx", title="Select Save File Location",filetypes=(("Word Document","*.docx"),))
        save_file =  filedialog.asksaveasfilename(initialfile="Photos.docx", defaultextension=".docx", title="Select Save File Location",filetypes=(("Word Document","*.docx"),))
    else: # Append Mode, Insert Mode - Load Document
        save_file = filedialog.askopenfilename(initialfile="Photos.docx", defaultextension=".docx", title="Select Existing File", filetypes=(("Word Document","*.docx"),))
    return save_file

def getExtensions():
    extensions = set()
    if jpg_extension.get():
        extensions.add('.JPG')
    if jpeg_extension.get():
        extensions.add('.JPEG')
    if png_extension.get():
        extensions.add('.PNG')
    return extensions

def folder_check():
    entry_lines = get_folder_entry()
    update_window()
    if entry_lines != ['']:
        for folder_entry in entry_lines:
            if not folder_entry:
                continue
            if not os.path.exists(folder_entry):
                folder_path_entry.configure(border_color="pink")
                status_label.configure(text="Folder(s) specified does not exist. Please re-enter.", fg_color="pink")
                return False
        status_label.configure(text="Configure options and click 'Submit' to begin collating.", fg_color="yellow")
        submit_button.configure(state="active")
        folder_path_entry.configure(fg_color=default_colour, border_color=default_border_colour)
        if not number_entry_checks():
            status_label.configure(text="Invalid number.", fg_color="pink")
        return True
    status_label.configure(text="Set folder location of photos.", fg_color="yellow")
    submit_button.configure(state="disabled")
    return False

def number_entry_check(number_widget):
    try:
        if not isfloat(number_widget.get()):
            raise ValueError
        else:
            if number_widget == start_count_entry:
                int(number_widget.get())
            number_widget.configure(fg_color=default_colour, border_color=default_border_colour)
            return True
    except ValueError:
        number_widget.configure(border_color="pink")
        return False

def number_entry_checks():
    valid_entry = True
    number_entry_widgets = [pic_width_entry, pic_height_entry, start_count_entry]
    for w in number_entry_widgets:
        if w.winfo_ismapped() and not number_entry_check(w):
            valid_entry = False
    return valid_entry

def validate_entries():
    valid_entries = number_entry_checks()
    if not caption_entry.get():
        valid_entries = False
    if not folder_check():
        valid_entries = False
    if not valid_entries:
        submit_button.configure(state="disabled")
    else:
        submit_button.configure(state="normal", hover=True, hover_color=default_hover_colour)
    return valid_entries

def isfloat(str):
    try:
        if not str:
            raise ValueError
        if float(str) <= 0:
            raise ValueError
    except ValueError:
        return False
    return True

@threaded_func
def update_label(text:str, *args, **kwargs):
    status_label.configure(text=text, *args, **kwargs)

def submit():
    if not validate_entries():
        return
    update_label(text="Processing..", fg_color="yellow")
    pic_extension = getExtensions()
    save_file = browse_save()
    if not save_file:
        status_label.configure(text="Configure options and click 'Submit' to begin collating.", fg_color="yellow")
        return
    
    try:
        folder_paths = [folder_path for folder_path in get_folder_entry() if folder_path]
        create_word_document(folder_paths, pic_width_entry.get(), pic_height_entry.get(), start_count_entry.get(), caption_entry.get(), pic_extension, save_file)
    except (ValueError, AttributeError, TypeError) as e:
        print("//ERROR:", type(e), e)
        status_label.configure(text="AN ERROR OCCURRED. PLEASE TRY AGAIN.", fg_color="red")
    except FileNotFoundError:
        status_label.configure(text="Folder location is not valid. Please reselect.", fg_color="red")

def get_folder_entry():
    return folder_path_entry.get('0.0','end').splitlines()

def update_window():
    global app_height
    lines = get_folder_entry()
    nlines = len(lines)
    nlines = nlines if nlines >= 0 else 0
    if nlines == 1 and not app_height:
        app_height = app.winfo_height()
        return
    # size_increment = font_size + 2
    size_increment = nlines * (font_size + 3)
    app.geometry(f"{app.winfo_width()}x{app_height+size_increment}")
    folder_path_entry.configure(height=widget_properties['height']+size_increment)
    # folder_path_entry.configure(height=folder_path_entry.cget('height')+size_increment)

def handle_menu(choice):
    print(choice, 'Selected')
    start_count_label.grid()
    start_count_entry.grid()
    if choice == menu_values[2]:
        start_count_label.configure(text="Insert Photo Before:")
    elif choice == menu_values[1]:
        start_count_label.grid_remove()
        start_count_entry.grid_remove()
    else:
        start_count_label.configure(text="Start Count:")
    folder_check()

# Create the main application window
app = ctk.CTk()
app.title("Photo Collator")
app.resizable(False, False)
app.grid_rowconfigure(4, minsize=40)

# append_mode = False
widget_row = 0
dw = ctk.CTkEntry(app)
widget_attributes = ['height', 'width', 'corner_radius', 'border_width', 'border_color', 'fg_color', 'text_color', 'font']
widget_properties = {k:dw.cget(k) for k in widget_attributes}
font_size = widget_properties['font'].cget('size')
app_height = False
app_mode = ctk.StringVar()
menu_values = ['Create Mode', 'Append Mode', 'Insert Mode']
app_mode.set(menu_values[0])
menubar = ctk.CTkOptionMenu(app, values=menu_values, variable=app_mode, command=handle_menu)
menubar.grid(row=widget_row, column=0, padx=10, pady=5, sticky='EW')

# Create and place widgets
widget_row += 1
folder_path_label = ctk.CTkLabel(app, text="Folder Location of Photos:").grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
folder_path_entry = ctk.CTkTextbox(app, wrap='none', **widget_properties)
folder_path_entry.grid(row=widget_row, column=1, padx=10, pady=5)
browse_button = ctk.CTkButton(app, text="Browse", command=browse_folder).grid(row=widget_row, column=2, padx=10, pady=5, sticky='EW')
folder_path_entry.bind("<FocusOut>", lambda event: validate_entries())
# folder_update_keys = ['Return', 'BackSpace', 'Up']
folder_path_entry.bind("<Key>", lambda event: folder_check())

widget_row += 1
pic_width_label = ctk.CTkLabel(app, text="Picture Width (cm):").grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
pic_width_entry = ctk.CTkEntry(app)
pic_width_entry.grid(row=widget_row, column=1, padx=10, pady=5, sticky='EW')
pic_width_entry.insert(0, "12.52")
pic_width_entry.bind("<FocusOut>", lambda event: validate_entries())

widget_row += 1
pic_height_label = ctk.CTkLabel(app, text="Picture Height (cm):").grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
pic_height_entry = ctk.CTkEntry(app)
pic_height_entry.grid(row=widget_row, column=1, padx=10, pady=5, sticky='EW')
pic_height_entry.insert(0, "8.3")
pic_height_entry.bind("<FocusOut>", lambda event: validate_entries())

widget_row += 1
start_count_label = ctk.CTkLabel(app, text="Start Count:")
start_count_label.grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
start_count_entry = ctk.CTkEntry(app)
start_count_entry.grid(row=widget_row, column=1, padx=10, pady=5, sticky='EW')
start_count_entry.insert(0, "1")
start_count_entry.bind("<FocusOut>", lambda event: validate_entries())

widget_row += 1
caption_label = ctk.CTkLabel(app, text="Default Caption:").grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
caption_entry = ctk.CTkEntry(app)
caption_entry.grid(row=widget_row, column=1, padx=10, pady=5, sticky='EW')
caption_entry.insert(0, "Photo")

widget_row += 1
pic_extension_label = ctk.CTkLabel(app, text="Picture Extension:").grid(row=widget_row, column=0, padx=10, pady=5, sticky='E')
jpg_extension = ctk.BooleanVar()
jpeg_extension = ctk.BooleanVar()
png_extension = ctk.BooleanVar()
jpg_checkbox = ctk.CTkCheckBox(app, text=".JPG", variable=jpg_extension, onvalue=True, offvalue=False)
jpg_checkbox.grid(row=widget_row, column=1, columnspan=1, padx=10, pady=5, sticky='W')
jpeg_checkbox = ctk.CTkCheckBox(app, text=".JPEG", variable=jpeg_extension, onvalue=True, offvalue=False)
jpeg_checkbox.grid(row=widget_row, column=1, columnspan=2, padx=80, pady=5, sticky='W')
png_checkbox = ctk.CTkCheckBox(app, text=".PNG", variable=png_extension, onvalue=True, offvalue=False)
png_checkbox.grid(row=widget_row, column=2, columnspan=2, padx=0, pady=5, sticky='W')
jpg_checkbox.select()
png_checkbox.select()

widget_row += 1
submit_button = ctk.CTkButton(app, text="Submit", command=submit, state="disabled")
submit_button.grid(row=widget_row, column=0, columnspan=3, padx=10, pady=10)

widget_row += 1
status_label = ctk.CTkLabel(app, text="Set folder location of photos.", corner_radius=5, padx=20, fg_color="yellow", text_color="black")
status_label.grid(row=widget_row, column=0, columnspan=3, padx=10, pady=5)
default_colour = folder_path_entry.cget("fg_color")
default_border_colour = folder_path_entry.cget("border_color")
default_hover_colour = submit_button.cget("hover_color")


# Start the Tkinter main loop
app.mainloop()
