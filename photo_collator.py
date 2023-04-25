#----------------------------------------------------------------------------------
#  CREATE ASPOSE DOC - FOR PAGE LAYOUT AND FORMAT
import aspose.words as aw

# -- ASPOSE DOCUMENT
adoc = aw.Document()
builder = aw.DocumentBuilder(adoc)

# -- PAGE SETUP
# Convert mm to pt
def aMm(dimension):
    return aw.ConvertUtil.millimeter_to_point(dimension)

# Margins (Size in mm)
vertical_margin = 17.5
horizontal_margin = 12.5

# Page Setup
builder.page_setup.orientation = aw.Orientation.LANDSCAPE
builder.page_setup.paper_size = aw.PaperSize.A4
builder.page_setup.top_margin = aMm(horizontal_margin)
builder.page_setup.bottom_margin = aMm(horizontal_margin)
builder.page_setup.left_margin = aMm(vertical_margin)
builder.page_setup.right_margin = aMm(vertical_margin)

# Save Aspose Document
adoc.save('test.docx')
print("Aspose Document Saved\nPreparing Python-Docx Document")
#----------------------------------------------------------------------------------

#----------------------------------------------------------------------------------
# CREATE PYTHON DOCX DOC - FOR IMAGES AND CAPTIONS
import os
from pathlib import Path
from docx import Document
from docx.shared import Cm, Pt
import natsort

# -- SETUP
fol = input("ENTER FOLDER OF IMAGES: ")
main_folder = Path(fol)
excluded = ['[excluded]', '.', '_']
extensions = ['JPG', 'PNG']
test = 0
number = 1

# -- DOCUMENT
doc = Document('test.docx')

# -- STYLES
style = doc.styles['Normal']
font = style.font
font.name = "Arial"
font.size = Pt(12)
# para_format = doc.paragraphs[-1].paragraph_format
# para_format.line_spacing = Pt(1.08)

# -- IMAGES
# Properties
# 0.0472p = 1cm
img_width = 12.52 # 12.52cm
img_height = 8.28 # 8.28cm

def test(test):
    for i in natsort.natsorted(test.glob('*.JPG'), alg=natsort.PATH):
        print(i)

def insert_photos(path):
    print(f'Adding Photos from {path.name}')
    for file in natsort.natsorted(path.glob('*.JPG'), alg=natsort.PATH):
        doc.paragraphs[-1].add_run().add_picture(str(file), width=Cm(img_width), height=Cm(img_height))
        doc.add_paragraph(f'Photo {number}')
        number += 1
        doc.add_paragraph('')
        if number % 2 == 1:
            doc.add_paragraph('')
    subfolders = [subfol for subfol in path.iterdir() if subfol.is_dir()]
    for subfol in subfolders:
        insert_photos(subfol)


# Search Photos
def find_photos(path):
    for file in path.glob('*'):
        if file.is_dir() and not file.name.lower().startswith(tuple(excluded)):                       # Check if folder
            #print(file)
            find_photos(file)                                         # Keep searching through subdirectories
        elif file.suffix.upper() in ['.'+ext for ext in extensions]:  # Check extension is a photo
            if test < 50:
                print(f"Adding Photo {number}: ", file)
                doc_photos(file)                                      # Pass photos only to document
                test += 1
                number += 1
                return
            else:
                return

# Add Photos
def doc_photos(photo, caption='Photo'):
    tabs = '\t'
    doc.paragraphs[-1].add_run().add_picture(str(photo), width=Cm(img_width), height=Cm(img_height))                           # Add to document
    if number % 2 == 1: # Odd
        doc.paragraphs[-1].add_run(tabs*2)
    else: # Even
        doc.add_paragraph(f"{caption} {number-1}{tabs*7}Photo {number}")
        doc.add_paragraph()
        if number % 4 != 0: # Middle of page - not end
            doc.add_paragraph()
    print("Done!")




# --- RUN PROGRAM

insert_photos(main_folder)
# find_photos(main_folder)

doc.save('test.docx')
print("Saved! Exiting..")