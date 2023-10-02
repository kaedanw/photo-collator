import tkinter as tk
from tkinter import filedialog
import os
from docx import Document # pip install python-docx | https://python-docx.readthedocs.io/en/latest/user/install.html
from docx.shared import Pt, Cm, Mm
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import natsort # pip install natsort | https://pypi.org/project/natsort/#installation

def create_word_document(folder_path, pic_width, pic_height, start_count, default_caption, pic_extensions, save_file):
    print("Creating Word Document..")
    
    # Get picture dimensions
    pic_width = float(pic_width)
    pic_height = float(pic_height)

    # Create a new Word document
    doc = Document()
    
    # Set landscape orientation and two columns layout
    section = doc.sections[0]
    section.page_width = Mm(297)
    section.page_height = Mm(210)
    section.orientation = WD_ORIENT.LANDSCAPE
    section.left_margin = Mm(17.5)
    section.right_margin = Mm(17.5)
    section.top_margin = Mm(12.5)
    section.bottom_margin = Mm(10.0)
    section.header_distance = Mm(12.7)
    section.footer_distance = Mm(12.7)
    sectPr = section._sectPr
    cols = sectPr.xpath('./w:cols')[0]
    cols.set(qn('w:num'),'2')
    
    # Set document style
    style = doc.styles['Normal']
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(8)
    paragraph_format.line_spacing = 1.1

    font = style.font
    font.name = 'Arial'
    font.size = Pt(12)

    # Get a list of files with the specified extension in the folder
    pic_files = [f for f in os.listdir(folder_path) if any(f.upper().endswith(e) for e in pic_extensions)]
    pic_files = natsort.natsorted(pic_files, alg=natsort.PATH)
    # Initialize photo counter
    count = int(start_count)
    
    # Iterate through the picture files and insert them into the document
    for pic_file in pic_files:     
        # Insert the image
        doc.add_paragraph()
        doc.paragraphs[-1].add_run().add_picture(os.path.join(folder_path, pic_file), width=Cm(pic_width), height=Cm(pic_height))
        
        # Add a caption below the image
        doc.add_paragraph()
        caption = f"{default_caption} {count}"
        doc.paragraphs[-1].add_run(caption)
        status_label.config(text=f"Inserting Photo {count}..")
        count += 1

    # Save the Word document
    status_label.config(text=f"Saving document..")
    doc.save(save_file)
    print("Word document created successfully.")

def browse_folder():
    folder_path = filedialog.askdirectory()
    folder_path_entry.delete(0, tk.END)
    folder_path_entry.insert(0, folder_path)
    folder_check()

def browse_save():
    desktop = os.path.normpath(os.path.expanduser("~/Desktop")).replace('\\','/')
    save_file =  filedialog.asksaveasfilename(initialdir=desktop, initialfile="Photos.docx", defaultextension=".docx", title="Select Save File Location",filetypes=(("Word Document","*.docx"),))
    return save_file

def getExtensions():
    extensions = set()
    if jpg_extension:
        extensions.add('.JPG')
    if jpeg_extension:
        extensions.add('.JPEG')
    if png_extension:
        extensions.add('.PNG')
    return extensions

def folder_check():
    if folder_path_entry.get():
        status_label.config(text="Configure options and click 'Submit' to begin collating.", bg="yellow")
        submit_button.config(state="active")
    else:
        status_label.config(text="Set folder location of photos.", bg="yellow")
        submit_button.config(state="disabled")

def validate_entries(folder_path, pic_width, pic_height, start_count, caption, pic_extension):
    if folder_path and pic_width and pic_height and start_count and caption and pic_extension:
        return True
    return False

def submit():
    folder_path = folder_path_entry.get()
    if not os.path.exists(folder_path):
        status_label.config(text="Folder specified does not exist. Please reselect.", bg="pink")
        return
    pic_width = pic_width_entry.get()
    pic_height = pic_height_entry.get()
    start_count = start_count_entry.get()
    caption = caption_entry.get()
    pic_extension = getExtensions()
    save_file = browse_save()

    status_label.config(text="Processing..", bg="yellow")
    validation = validate_entries(folder_path, pic_width, pic_height, start_count, caption, pic_extension)
    if not validation:
        status_label.config(text="Please fill out all fields.", bg="pink")
        return
    if not save_file:
        status_label.config(text="Configure options and click 'Submit' to begin collating.", bg="yellow")
        return
    
    try:
        create_word_document(folder_path, pic_width, pic_height, start_count, caption, pic_extension, save_file)
        status_label.config(text="FINISHED SUCCESSFULLY", bg="lightgreen")
    except (ValueError, AttributeError):
        status_label.config(text="AN ERROR OCCURRED. PLEASE TRY AGAIN.", bg="red")
    except FileNotFoundError:
        status_label.config(text="Folder location is not valid. Please reselect.", bg="red")

# Create the main application window
app = tk.Tk()
app.title("Collate Document Photos App")

# Create and place widgets
folder_path_label = tk.Label(app, text="Folder Location of Photos:").grid(row=0, column=0, padx=10, pady=5, sticky='E')

folder_path_entry = tk.Entry(app)
folder_path_entry.grid(row=0, column=1, padx=10, pady=5)
browse_button = tk.Button(app, text="Browse", command=browse_folder).grid(row=0, column=2, padx=10, pady=5, sticky='EW')
# app.bind("<FocusOut>", lambda _: folder_focus())
# app.bind("<FocusIn>", lambda _: folder_focus())


pic_width_label = tk.Label(app, text="Picture Width (cm):").grid(row=1, column=0, padx=10, pady=5, sticky='E')
pic_width_entry = tk.Entry(app)
pic_width_entry.grid(row=1, column=1, padx=10, pady=5, sticky='EW')
pic_width_entry.insert(0, "12.52")

pic_height_label = tk.Label(app, text="Picture Height (cm):").grid(row=2, column=0, padx=10, pady=5, sticky='E')
pic_height_entry = tk.Entry(app)
pic_height_entry.grid(row=2, column=1, padx=10, pady=5, sticky='EW')
pic_height_entry.insert(0, "8.3")

start_count_label = tk.Label(app, text="Start Count:").grid(row=3, column=0, padx=10, pady=5, sticky='E')
start_count_entry = tk.Entry(app)
start_count_entry.grid(row=3, column=1, padx=10, pady=5, sticky='EW')
start_count_entry.insert(0, "1")

caption_label = tk.Label(app, text="Default Caption:").grid(row=4, column=0, padx=10, pady=5, sticky='E')
caption_entry = tk.Entry(app)
caption_entry.grid(row=4, column=1, padx=10, pady=5, sticky='EW')
caption_entry.insert(0, "Photo")

pic_extension_label = tk.Label(app, text="Picture Extension:").grid(row=5, column=0, padx=10, pady=5, sticky='E')
# pic_extension_entry = tk.Entry(app)
# pic_extension_entry.insert(0, "JPG")
# pic_extension_entry.grid(row=5, column=1, padx=10, pady=5, sticky='EW')
jpg_extension = tk.BooleanVar()
jpeg_extension = tk.BooleanVar()
png_extension = tk.BooleanVar()
jpg_checkbox = tk.Checkbutton(app, text=".JPG", variable=jpg_extension, onvalue=True, offvalue=False)
jpg_checkbox.grid(row=5, column=1, padx=5, pady=5, sticky='W')
jpeg_checkbox = tk.Checkbutton(app, text=".JPEG", variable=jpeg_extension, onvalue=True, offvalue=False)
jpeg_checkbox.grid(row=5, column=1, padx=15, pady=5, sticky='E')
png_checkbox = tk.Checkbutton(app, text=".PNG", variable=png_extension, onvalue=True, offvalue=False)
png_checkbox.grid(row=5, column=2, padx=0, pady=5, sticky='W')
jpg_checkbox.select()
png_checkbox.select()

submit_button = tk.Button(app, text="Submit", command=submit, state="disabled")
submit_button.grid(row=7, column=0, columnspan=3, padx=10, pady=10)

status_label = tk.Label(app, text="Set folder location of photos.", bg="yellow")
status_label.grid(row=8, column=0, columnspan=3, padx=10, pady=5)


# Start the Tkinter main loop
app.mainloop()
